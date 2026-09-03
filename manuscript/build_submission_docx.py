#!/usr/bin/env python3
"""Build the Word submission artifact from the journal-facing manuscript.

The Markdown manuscript uses raw LaTeX figure blocks for the PDF build.  This
builder converts those blocks to portable Markdown images for DOCX, lets
Pandoc resolve the conventional citations, and applies the explicit
standard_business_brief style tokens required for the Word deliverable.
"""

from __future__ import annotations

import argparse
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FIGURE_RE = re.compile(
    r"\\begin\{figure\}\[p\]\s*"
    r"\\centering\s*"
    r"\\includegraphics\[[^\]]+\]\{(?P<path>[^}]+)\}\s*"
    r"\\par\\textbf\{(?P<label>Fig\.\s*\d+)\}\s+(?P<caption>[^\n]+)\s*"
    r"\\end\{figure\}",
    re.MULTILINE,
)


def set_font(run, name="Calibri", size=11, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                set_spacing(paragraph, after=2, line=1.0)
                for run in paragraph.runs:
                    set_font(run, size=9)
            if row_index == 0:
                shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
                if shd is None:
                    shd = OxmlElement("w:shd")
                    cell._tc.get_or_add_tcPr().append(shd)
                shd.set(qn("w:fill"), "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        if row_index == 0:
            tr_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def prepare_markdown(source: str) -> str:
    def replace(match):
        path = match.group("path")
        label = match.group("label").replace("  ", " ")
        caption = match.group("caption").strip()
        return f"![{label} {caption}]({path})"

    return FIGURE_RE.sub(replace, source)


def style_document(path: Path):
    doc = Document(path)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for index, paragraph in enumerate(doc.paragraphs):
        set_spacing(paragraph)
        if index == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_spacing(paragraph, after=10, line=1.0)
            for run in paragraph.runs:
                set_font(run, size=20, bold=True, color="0B2545")
        elif paragraph.text.startswith(("Article type:", "Authors and affiliations:", "Corresponding author:", "ORCID iDs:")):
            set_spacing(paragraph, after=3, line=1.0)
            for run in paragraph.runs:
                set_font(run, size=10, color="555555")
        elif paragraph.style.name.startswith("List"):
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            set_spacing(paragraph, after=8, line=1.167)
        elif paragraph.text.startswith("Fig. "):
            set_spacing(paragraph, before=4, after=8, line=1.0)
            for run in paragraph.runs:
                set_font(run, size=10, color="555555")

    widths = [1800, 2800, 2400, 2360]
    for table in doc.tables:
        set_table_geometry(table, widths)

    # Named override: journal_figure_max_height = 7.0 in.  The source figures
    # are intentionally tall diagrams; capping their height keeps the complete
    # image and caption on a Letter page without distorting the aspect ratio.
    for shape in doc.inline_shapes:
        max_height = Inches(7.0)
        if shape.height > max_height:
            ratio = shape.width / shape.height
            shape.height = max_height
            shape.width = int(max_height * ratio)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.text = "Artemisia terpene review · submission manuscript"
    set_spacing(header, after=0, line=1.0)
    for run in header.runs:
        set_font(run, size=8, color="777777")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = ""
    run = footer.add_run("Page ")
    set_font(run, size=8, color="777777")
    add_page_field(footer)
    for run in footer.runs:
        set_font(run, size=8, color="777777")

    doc.save(path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, default=Path("article-submission.md"))
    parser.add_argument("--output", type=Path, default=Path("artemisia-terpene-review-submission.docx"))
    args = parser.parse_args()
    source = args.input.resolve()
    args.output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile("w", suffix=".md", encoding="utf-8", delete=False) as handle:
        handle.write(prepare_markdown(source.read_text(encoding="utf-8")))
        temp_input = Path(handle.name)
    try:
        subprocess.run(
            [
                "pandoc",
                str(temp_input),
                "--from",
                "markdown",
                "--citeproc",
                "--bibliography",
                str(source.parent / "references.bib"),
                "--resource-path",
                str(source.parent),
                "-o",
                str(args.output.resolve()),
            ],
            check=True,
        )
    finally:
        temp_input.unlink(missing_ok=True)
    style_document(args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
